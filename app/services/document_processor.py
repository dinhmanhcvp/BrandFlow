import os
import re
from datetime import datetime
from typing import List
from dotenv import load_dotenv

load_dotenv()

from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_chroma import Chroma
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain_core.documents import Document

class DocumentIngestor:
    def __init__(self, persist_directory: str = "./chroma_db", tenant_id: str = "default"):
        self.persist_directory = persist_directory
        self.tenant_id = tenant_id
        self.embeddings = GoogleGenerativeAIEmbeddings(model="models/gemini-embedding-001")
        
    def ingest_file(self, file_path: str, force_ai: bool = False) -> str:
        """
        Đọc file bằng cơ chế Smart Parser (PDF -> LlamaParse AI).
        """
        print(f"📄 [Ingestion] Đang quét nội dung: {file_path}...")
        ext = file_path.lower().split('.')[-1]
        raw_text = ""
        try:
            if ext in ["txt", "md", "csv"]:
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    raw_text = f.read()
            elif ext == "docx":
                import docx
                doc = docx.Document(file_path)
                raw_text = "\n".join([para.text for para in doc.paragraphs])
            elif ext == "pdf":
                raw_text = ""
                api_key = os.getenv("LLAMA_CLOUD_API_KEY")
                if api_key:
                    print(f"   🚀 Phát hiện file PDF! Gọi ngay AI LlamaParse Cloud...")
                    try:
                        from llama_parse import LlamaParse
                        parser = LlamaParse(
                            api_key=api_key,
                            result_type="markdown",
                            verbose=False
                        )
                        parsed_docs = parser.load_data(file_path)
                        raw_text = "\n".join([doc.text for doc in parsed_docs])
                    except Exception as parse_e:
                        print(f"   ⚠️ Lỗi khi gọi LlamaParse: {parse_e}")
                
                if not raw_text.strip():
                    import pdfplumber
                    try:
                        with pdfplumber.open(file_path) as pdf:
                            for page in pdf.pages:
                                extracted = page.extract_text()
                                if extracted:
                                    raw_text += extracted + "\n"
                    except Exception:
                        raw_text = ""
            else:
                print(f"   ⚠️ [Bỏ qua] Định dạng {ext} không hỗ trợ.")
                
            return raw_text
        except Exception as e:
            print(f"🔴 [Ingestion Lỗi]: {e}")
            raise ValueError(f"Ingestion failed: {e}")

    def ingest_url(self, url: str) -> str:
        """Đọc nội dung text từ đường dẫn URL."""
        try:
            from langchain_community.document_loaders import WebBaseLoader
            loader = WebBaseLoader(url)
            docs = loader.load()
            return "\n".join([doc.page_content for doc in docs])
        except Exception as e:
            raise ValueError(f"Ingestion URL failed: {e}")

    def clean_text(self, text: str) -> str:
        """Làm sạch văn bản cơ bản."""
        if not text: return ""
        text = text.replace("\x00", "") 
        text = re.sub(r'[ \t]+', ' ', text)
        text = re.sub(r'\n{3,}', '\n\n', text)
        return text.strip()

    def semantic_chunking(self, text: str) -> List[str]:
        """Cắt văn bản thành các fragments tối ưu cho RAG."""
        print("✂️  [Chunking] Đang phân mảnh tài liệu...")
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=150,
            separators=["\n\n", "\n", ".", " ", ""]
        )
        chunks = splitter.split_text(text)
        print(f"   ✅ Đã phân thành {len(chunks)} chunks.")
        return chunks

    def process_and_store_text(self, raw_text: str, filename: str, category: str):
        """Pipeline: Cleansing -> Chunking -> VectorDB."""
        current_date = datetime.now().isoformat()
        try:
            cleaned_text = self.clean_text(raw_text)
            chunks = self.semantic_chunking(cleaned_text)
            if not chunks: return
            
            documents = [
                Document(
                    page_content=chunk,
                    metadata={"source": filename, "category": category, "date": current_date, "index": i}
                ) for i, chunk in enumerate(chunks)
            ]
            
            print(f"💾 [Storage] Đang lưu {len(documents)} bản ghi vào ChromaDB...")
            collection_name = f"brandflow_memory_{self.tenant_id}"
            vectorstore = Chroma(
                collection_name=collection_name,
                embedding_function=self.embeddings,
                persist_directory=self.persist_directory
            )
            vectorstore.add_documents(documents)
            print(f"🎉 Đã lưu thành công [{filename}] vào bộ não.")
        except Exception as e:
            print(f"🔴 [Storage Lỗi]: {e}")
