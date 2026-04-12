"""
=============================================================================
BrandFlow Strategy Engine - memory_rag.py (v7 — Long-Term Memory)
=============================================================================
Quản lý bộ nhớ dài hạn cho hệ thống đa Agent.

Chức năng:
  1. ChromaDB: Vector Database lưu trữ quy tắc công ty & bài học kinh nghiệm.
  2. Learner Agent: Trích xuất quy tắc từ kế hoạch bị từ chối.
  3. RAG Retrieval: Truy xuất quy tắc liên quan trước khi MasterPlanner lập kế hoạch.

Chạy 100% Local: OllamaEmbeddings (nomic-embed-text) + ChatOllama (llama3.2).
=============================================================================
"""

import sys
if sys.stdout.encoding and sys.stdout.encoding.lower() != 'utf-8':
    sys.stdout.reconfigure(encoding='utf-8')

import json
import io
from typing import List
from pydantic import BaseModel, Field

import pdfplumber
import docx
from fastapi import UploadFile

# ---- LangChain imports ----
from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
from langchain_chroma import Chroma
from langchain_core.output_parsers import JsonOutputParser
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.documents import Document

# =============================================================================
# 1. KHỞI TẠO EMBEDDINGS & VECTOR STORE
# =============================================================================

CHROMA_PERSIST_DIR = "./chroma_db"

def get_embeddings() -> GoogleGenerativeAIEmbeddings:
    """Khởi tạo Embedding model (Gemini)."""
    return GoogleGenerativeAIEmbeddings(model="models/gemini-embedding-001")


def get_vectorstore(tenant_id: str = "default") -> Chroma:
    """Khởi tạo hoặc mở lại ChromaDB tại thư mục ./chroma_db."""
    return Chroma(
        collection_name=f"brandflow_memory_{tenant_id}",
        embedding_function=get_embeddings(),
        persist_directory=CHROMA_PERSIST_DIR,
    )


# =============================================================================
# 2. LEARNER AGENT — Trích xuất bài học từ kế hoạch bị từ chối
# =============================================================================

class LearnedRule(BaseModel):
    """Schema cho quy tắc được trích xuất từ sai lầm."""
    rule_summary: str = Field(description="Tóm tắt quy tắc ngắn gọn (1-2 câu)")
    keywords: List[str] = Field(description="Danh sách từ khóa để tìm kiếm (3-5 từ)")


learner_parser = JsonOutputParser(pydantic_object=LearnedRule)

learner_prompt = ChatPromptTemplate.from_messages([
    (
        "system",
        """Bạn là chuyên gia phân tích thất bại. Đọc bản kế hoạch bị từ chối và lời phê bình của Giám đốc.
Nhiệm vụ: Trích xuất MỘT quy tắc tổng quát mà công ty nên tuân thủ trong tương lai.

Quy tắc phải:
- Ngắn gọn (1-2 câu).
- Mang tính tổng quát, áp dụng được cho nhiều chiến dịch, không chỉ riêng chiến dịch này.
- Kèm theo 3-5 từ khóa liên quan.

CHỈ TRẢ VỀ CHUỖI JSON HỢP LỆ. KHÔNG CÓ BẤT KỲ VĂN BẢN NÀO BÊN NGOÀI.

{format_instructions}"""
    ),
    (
        "human",
        """Bản kế hoạch bị từ chối:
{rejected_plan}

Lời phê bình của Giám đốc (Human Feedback):
{human_feedback}

Hãy trích xuất quy tắc rút kinh nghiệm."""
    ),
])


def extract_and_save_rule(human_feedback: str, rejected_plan: str, tenant_id: str = "default") -> str:
    """
    Cho Learner Agent đọc kế hoạch bị chê + feedback của người dùng.
    Trích xuất quy tắc tổng quát và lưu vào ChromaDB.

    Returns:
        rule_summary (str) đã lưu thành công, hoặc chuỗi lỗi.
    """
    try:
        llm = ChatGoogleGenerativeAI(model="gemini-2.5-flash", temperature=0.0)
        chain = (
            learner_prompt.partial(
                format_instructions=learner_parser.get_format_instructions()
            )
            | llm
            | learner_parser
        )

        result = chain.invoke({
            "rejected_plan": rejected_plan,
            "human_feedback": human_feedback,
        })

        rule_summary = result.get("rule_summary", "Không trích xuất được quy tắc.")
        keywords = result.get("keywords", [])

        # Lưu vào ChromaDB
        vectorstore = get_vectorstore(tenant_id)
        doc = Document(
            page_content=rule_summary,
            metadata={
                "type": "lesson",
                "keywords": json.dumps(keywords, ensure_ascii=False),
                "source_feedback": human_feedback[:200],
            }
        )
        vectorstore.add_documents([doc])

        print(f"   📝 [Learner] Đã lưu quy tắc mới vào bộ nhớ dài hạn:")
        print(f"      → {rule_summary}")
        print(f"      → Keywords: {', '.join(keywords)}")

        return rule_summary

    except Exception as e:
        print(f"   🔴 [Learner] Lỗi khi trích xuất quy tắc: {e}")
        return f"Lỗi: {str(e)[:100]}"


# =============================================================================
# 3. FILE PARSING & BRAND DNA EXTRACTION
# =============================================================================

async def parse_file_content(file: UploadFile) -> str:
    """Đọc nội dung thô từ file PDF, DOCX, hoặc TXT tải lên."""
    content = await file.read()
    filename = file.filename.lower()
    text = ""

    try:
        if filename.endswith(".pdf"):
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                for page in pdf.pages:
                    extracted = page.extract_text()
                    if extracted:
                        text += extracted + "\n"
        elif filename.endswith(".docx"):
            doc = docx.Document(io.BytesIO(content))
            for para in doc.paragraphs:
                text += para.text + "\n"
        elif filename.endswith(".txt"):
            text = content.decode("utf-8")
        else:
            raise ValueError("Định dạng file không được hỗ trợ. Chỉ nhận PDF, DOCX, TXT.")
    except Exception as e:
        print(f"🔴 [File Parser] Lỗi khi đọc file {file.filename}: {e}")
        raise ValueError(f"Không thể đọc nội dung file: {str(e)}")

    return text.strip()


class BrandDNA(BaseModel):
    """Schema cho phân tích Brand DNA từ file khách hàng tải lên."""
    core_usps: List[str] = Field(description="3-5 điểm bán hàng độc nhất (Unique Selling Points)")
    target_audience_insights: List[str] = Field(description="Chân dung và Insights khách hàng mục tiêu")
    tone_of_voice: str = Field(description="Giọng điệu thương hiệu (VD: Hiện đại, hài hước, chuyên gia...)")
    strict_rules: List[str] = Field(description="Các quy tắc DOs và DON'Ts quan trọng cho Marketing")


dna_parser = JsonOutputParser(pydantic_object=BrandDNA)

dna_prompt = ChatPromptTemplate.from_messages([
    (
        "system",
        """Bạn là Giám đốc Chiến lược Thương hiệu (Brand Strategist) bậc thầy. Đọc tài liệu của khách hàng tải lên.
Nhiệm vụ: Trích xuất các thông tin cốt lõi (Brand DNA) để làm nền tảng cho AI MasterPlanner lập kế hoạch sau này.

CHỈ TRẢ VỀ CHUỖI JSON HỢP LỆ. KHÔNG CÓ BẤT KỲ VĂN BẢN NÀO BÊN NGOÀI.

{format_instructions}"""
    ),
    (
        "human",
        """Đây là toàn bộ nội dung tài liệu của doanh nghiệp:
{document_content}

Hãy trích xuất Brand DNA ngay lập tức."""
    ),
])


def analyze_and_extract_dna(document_content: str, tenant_id: str = "default") -> dict:
    """
    Sử dụng LLM để đọc text thô, trích xuất cấu trúc BrandDNA JSON.
    Sau đó tự động lưu các strict_rules vào ChromaDB.
    """
    if len(document_content) < 50:
        raise ValueError("Nội dung file quá ngắn hoặc không có text để phân tích.")

    # Cắt bớt text nếu quá dài cho Context Window của LLM cục bộ (giả định max 15000 ký tự cho an toàn)
    safe_content = document_content[:15000]

    try:
        llm = ChatGoogleGenerativeAI(model="gemini-2.5-flash", temperature=0.1)
        chain = (
            dna_prompt.partial(
                format_instructions=dna_parser.get_format_instructions()
            )
            | llm
            | dna_parser
        )

        print(f"🧠 [Brand DNA] Bắt đầu phân tích tài liệu ({len(safe_content)} ký tự)...")
        result = chain.invoke({"document_content": safe_content})

        # Lưu các strict_rules vào ChromaDB để AI nhớ mãi mãi
        vectorstore = get_vectorstore(tenant_id)
        strict_rules = result.get("strict_rules", [])
        if strict_rules:
            docs = [
                Document(
                    page_content=f"[BRAND RULE]: {rule}",
                    metadata={"type": "brand_dna_rule"}
                )
                for rule in strict_rules
            ]
            vectorstore.add_documents(docs)
            print(f"   ✅ [Brand DNA] Đã lưu {len(docs)} quy tắc bắt buộc vào bộ nhớ dài hạn.")

        return {
            "status": "success",
            "message": "Phân tích tài liệu thành công.",
            "data": result
        }

    except Exception as e:
        print(f"🔴 [Brand DNA] Lỗi khi LLM phân tích: {e}")
        return {"status": "error", "message": str(e)}


# =============================================================================
# 4. ONBOARDING MODULE — Khởi tạo dữ liệu (Cold Start)
# =============================================================================

def inject_industry_presets(industry_name: str, tenant_id: str = "default") -> dict:
    """Nạp bộ quy chuẩn ngành có sẵn vào ChromaDB."""
    presets = {
        "F&B": [
            "Luôn nhấn mạnh vào hình ảnh món ăn chân thực và hấp dẫn.",
            "Khuyến khích sử dụng các bài review thực tế từ khách hàng.",
            "Tuyệt đối không sử dụng các từ ngữ mang tính y khoa hoặc chữa bệnh."
        ],
        "Spa_Beauty": [
            "Bắt buộc phải có hình ảnh Before/After khi quảng cáo liệu trình.",
            "Không được cam kết chữa khỏi 100% các vấn đề về da/dáng.",
            "Sử dụng tone giọng nhẹ nhàng, thể hiện sự chuyên nghiệp và thấu hiểu của chuyên gia."
        ],
        "B2B_Tech": [
            "Tập trung truyền thông trên kênh LinkedIn với nội dung chuyên sâu.",
            "Nhấn mạnh vào tỷ lệ hoàn vốn (ROI) và các tính năng kỹ thuật nổi bật.",
            "Văn phong chuyên nghiệp, số liệu rõ ràng, tránh dùng từ lóng hoặc quá cảm xúc."
        ]
    }
    
    if industry_name not in presets:
        return {"status": "error", "message": f"Ngành '{industry_name}' không được hỗ trợ."}
        
    rules = presets[industry_name]
    try:
        vectorstore = get_vectorstore(tenant_id)
        docs = [
            Document(page_content=rule, metadata={"type": "preset", "industry": industry_name})
            for rule in rules
        ]
        vectorstore.add_documents(docs)
        print(f"   ✅ [Onboarding] Đã nạp {len(rules)} quy chuẩn cho ngành {industry_name}")
        return {"status": "success", "message": f"Đã nạp bộ quy chuẩn ngành {industry_name} thành công."}
    except Exception as e:
         print(f"   🔴 [Onboarding] Lỗi khi nạp rules: {e}")
         return {"status": "error", "message": str(e)}

def generate_guideline_from_qa(qa_pairs: dict, tenant_id: str = "default") -> dict:
    """Sinh ra 3 quy tắc marketing ngắn gọn từ câu trả lời phỏng vấn."""
    
    qa_text = "\n".join([f"Hỏi: {q}\nĐáp: {a}" for q, a in qa_pairs.items()])
    
    prompt = ChatPromptTemplate.from_messages([
        ("system", "Bạn là chuyên gia chiến lược thương hiệu. Dựa vào các câu trả lời phỏng vấn sau của chủ doanh nghiệp, hãy đúc kết ra đúng 3 quy tắc marketing (Brand Guidelines) ngắn gọn nhất để AI tuân thủ. Chỉ trả về các quy tắc, mỗi quy tắc bắt đầu bằng một dấu gạch ngang '-', không kèm giải thích hay bất kỳ văn bản nào khác."),
        ("human", "Câu trả lời phỏng vấn:\n{qa_text}")
    ])
    
    try:
        llm = ChatGoogleGenerativeAI(model="gemini-2.5-flash", temperature=0.2)
        chain = prompt | llm
        
        response = chain.invoke({"qa_text": qa_text})
        
        # Parse output
        rules = [line.strip().lstrip('- ') for line in response.content.split('\n') if line.strip().startswith('-')]
        
        if not rules:
             # Fallback incase format isn't strictly followed
             rules = [rule.strip() for rule in response.content.split('\n') if rule.strip()]
             
        # Optional constraint for 3 rules
        rules = rules[:3]
             
        vectorstore = get_vectorstore(tenant_id)
        docs = [
            Document(page_content=rule, metadata={"type": "onboarding"})
            for rule in rules
        ]
        vectorstore.add_documents(docs)
        print(f"   ✅ [Onboarding] Đã tạo và lưu {len(rules)} quy tắc từ phỏng vấn.")
        return {"status": "success", "message": "Đã khởi tạo sổ tay thương hiệu từ phỏng vấn.", "rules": rules}
        
    except Exception as e:
        print(f"   🔴 [Onboarding] Lỗi khi tạo guideline từ QA: {e}")
        return {"status": "error", "message": str(e)}

# =============================================================================
# 4. RAG RETRIEVAL — Truy xuất quy tắc liên quan
# =============================================================================

def get_relevant_guidelines(goal: str, top_k: int = 3, tenant_id: str = "default") -> str:
    """
    Query ChromaDB bằng mục tiêu chiến dịch, lấy ra top-k quy tắc liên quan nhất.

    Args:
        goal: Mục tiêu chiến dịch marketing.
        top_k: Số lượng quy tắc tối đa cần lấy.

    Returns:
        Chuỗi string gộp các quy tắc, hoặc "" nếu DB trống / lỗi.
    """
    try:
        vectorstore = get_vectorstore(tenant_id)

        # Kiểm tra xem DB có dữ liệu không
        collection = vectorstore._collection
        if collection.count() == 0:
            return ""

        results = vectorstore.similarity_search(goal, k=top_k)

        if not results:
            return ""

        guidelines = []
        for i, doc in enumerate(results, 1):
            guidelines.append(f"  {i}. {doc.page_content}")

        return "\n".join(guidelines)

    except Exception as e:
        print(f"   ⚠️ [RAG] Lỗi khi truy xuất bộ nhớ: {e}")
        return ""


def add_manual_guideline(text: str, guideline_type: str = "company_rule", tenant_id: str = "default") -> None:
    """Thêm một quy tắc thủ công vào ChromaDB."""
    try:
        vectorstore = get_vectorstore(tenant_id)
        doc = Document(
            page_content=text,
            metadata={"type": guideline_type}
        )
        vectorstore.add_documents([doc])
        print(f"   ✅ [DB] Đã lưu quy tắc: '{text[:80]}...'")
    except Exception as e:
        print(f"   🔴 [DB] Lỗi khi lưu: {e}")


# =============================================================================
# 4. TEST — Thử nghiệm trực tiếp
# =============================================================================

if __name__ == "__main__":
    print("\n" + "═" * 70)
    print("🧪 [TEST] memory_rag.py — Vector Database & RAG Retrieval")
    print("═" * 70)

    # (1) Thêm tay một Document mẫu
    print("\n📥 Đang thêm quy tắc mẫu vào ChromaDB...")
    add_manual_guideline(
        "Luật công ty: Ngân sách dưới 20 triệu không được thuê KOL.",
        guideline_type="company_rule"
    )
    add_manual_guideline(
        "Bài học: Không nên bỏ hết ngân sách vào một kênh duy nhất. Phải đa dạng ít nhất 3 kênh.",
        guideline_type="lesson"
    )

    # (2) Gọi thử hàm get_relevant_guidelines
    print("\n🔍 Đang truy vấn ChromaDB với mục tiêu: 'Chiến dịch tiết kiệm'...")
    guidelines = get_relevant_guidelines("Chiến dịch tiết kiệm")

    if guidelines:
        print(f"\n   📋 Kết quả tìm kiếm:")
        print(guidelines)
    else:
        print("   ⚠️ Không tìm thấy quy tắc nào!")

    print("\n" + "═" * 70)
    print("✅ Test hoàn tất!")
    print("═" * 70)
